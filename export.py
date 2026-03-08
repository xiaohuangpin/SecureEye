from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os,logging,base64,io

def export_to_word(data_list:list[dict[str,str]]) -> str:
    
    doc = Document()
    
    title = doc.add_heading('安全隐患整改单', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'  
    table.autofit = False
    
    
    col_widths = [Inches(1.5), Inches(1.0), Inches(1.5)]
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width
    
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '隐患图片'
    hdr_cells[1].text = '隐患描述'
    hdr_cells[2].text = '整改图片'
    
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    for item in data_list:
        
        row_cells = table.add_row().cells
        
        
        img_cell = row_cells[0]
        img:str = item['image']

        if img.startswith('data:image'):
            img = img.split(',')[1]
        
        image_bytes:bytes = base64.b64decode(img)

        image_stream:io.BytesIO = io.BytesIO(image_bytes) #: BinaryIO
        
        """
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
            img.save(tmpfile, format='PNG')
            tmp_path = tmpfile.name
        """

            
        paragraph = img_cell.paragraphs[0] if img_cell.paragraphs else img_cell.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
           
        run.add_picture(image_stream, width=Inches(1.2))
        
        label_cell = row_cells[1]
        label_cell.text = item['label']
        label_paragraph = label_cell.paragraphs[0] if label_cell.paragraphs else label_cell.add_paragraph()
        label_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        new_img_cell = row_cells[2]
        new_img_paragraph = new_img_cell.paragraphs[0] if new_img_cell.paragraphs else new_img_cell.add_paragraph()
        new_img_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    current_dir = os.path.dirname(os.path.abspath(__file__))
    output_path = os.path.join(current_dir, "image_table.docx")
    
    doc.save(output_path)
    logging.info(f"{output_path}文件保存成功")
    return output_path

if __name__ == "__main__":
    from agent import secure_cv
    import asyncio
    api_key="40d63b64ed374134a0b3b24c6e3963b0.jZy5E9Q37TUXmSXX"
    base_url="https://open.bigmodel.cn/api/paas/v4/"
    model_name = "GLM-4V-Flash"
    example = ["未固定的高空作业平台","裸露的带电电缆"]
    image_path = "image/14.jpg"
    ima_list = ['image/6.jpg','image/6.jpg']
    
    async def test() -> None:
        example_data = await secure_cv(api_key,base_url,model_name,example,ima_list)
        export_to_word(example_data)

    asyncio.run(test())

    #api = MultClient(api_key,base_url,model_name,example)
    #example_data = api.batch_infer(ima_list,True)
    #export_to_word(example_data)
    
    """
    ####启动文件
    import multiprocessing
    DOCX_FILE_PATH = "image_table.docx"
    file_process = multiprocessing.Process(
        target=open_file_with_default_app,
        args=(DOCX_FILE_PATH,),
        name="DocxOpenProcess"  # 进程命名，方便任务管理器识别
    )
    file_process.start()
    file_process.join()
    """
    