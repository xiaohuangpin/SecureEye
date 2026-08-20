from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os,logging,base64,io
from PIL import Image
import xlsxwriter

def export_to_word(data_list: list[dict[str, str | Image.Image]]) -> str:
    # 创建 docx对象
    doc:Document = Document()

    # 1. 添加标题
    title = doc.add_heading('安全隐患整改单', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 2. 创建表格并设置列宽
    table = doc.add_table(rows=1, cols=3, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    col_widths = [Inches(1.5), Inches(1.0), Inches(1.5)]
    for row in table.rows:
        for idx, width in enumerate(col_widths):
            row.cells[idx].width = width

    # 3. 设置表头
    headers = ['隐患图片', '隐患描述', '整改图片']
    hdr_cells = table.rows[0].cells
    for cell, header in zip(hdr_cells, headers):
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # 4. 填充数据行
    for item in data_list:
        row_cells = table.add_row().cells
        
       
        img_data = item['image']
        image_stream = io.BytesIO()
        if isinstance(img_data, Image.Image):
            img_data.save(image_stream, format='PNG')
        elif isinstance(img_data, str) and img_data.startswith('data:image'):
            image_stream.write(base64.b64decode(img_data.split(',')[1]))
        else:
            raise TypeError(f"不支持的图片数据类型: {type(img_data)}")

        # 重置流指针到开头，供 docx 读取
        image_stream.seek(0)
        
        # 插入图片
        img_cell = row_cells[0]
        img_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_cell.paragraphs[0].add_run().add_picture(image_stream, width=Inches(1.2))
        
        # 插入文本描述
        label_cell = row_cells[1]
        label_cell.text = item['label']
        label_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 设置所有单元格垂直居中
        for cell in row_cells:
            cell.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 5. 保存文件
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "image_table.docx")
    doc.save(output_path) # 保存docx文件
    logging.info(f"{output_path}文件保存成功")
    return output_path

def export_to_excel(data_list: list[dict[str, str | Image.Image]]) -> str:
    output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "image_table.xlsx")
    
    # 1. 创建工作簿和工作表
    workbook = xlsxwriter.Workbook(output_path)
    worksheet = workbook.add_worksheet('安全隐患整改单')

    # 2. 定义样式
    center_format = workbook.add_format({'align': 'center', 'valign': 'vcenter'})
    header_format = workbook.add_format({'align': 'center', 'valign': 'vcenter', 'bold': True})

    # 3. 写入表头
    headers = ['隐患图片', '隐患描述', '整改图片']
    for col_num, header in enumerate(headers):
        worksheet.write(0, col_num, header, header_format)

    # 4. 设置列宽 (Excel列宽单位近似为字符数)
    worksheet.set_column('A:A', 20)  # 隐患图片 (约 115px 宽)
    worksheet.set_column('B:B', 25)  # 隐患描述
    worksheet.set_column('C:C', 20)  # 整改图片

    # 5. 填充数据行
    for row_idx, item in enumerate(data_list, start=1):
        img_data = item['image']
        
        # --- 核心优化：统一转换为 PIL 对象并固定宽度，保证表格整齐 ---
        if isinstance(img_data, Image.Image):
            pil_img = img_data
        elif isinstance(img_data, str) and img_data.startswith('data:image'):
            pil_img = Image.open(io.BytesIO(base64.b64decode(img_data.split(',')[1])))
        elif isinstance(img_data, bytes):
            pil_img = Image.open(io.BytesIO(img_data))
        else:
            raise TypeError(f"不支持的图片数据类型: {type(img_data)}")
        
        # 统一调整宽度为 115 像素 (约 1.2 英寸)，高度按比例自适应
        target_width = 115
        ratio = target_width / pil_img.width
        target_height = int(pil_img.height * ratio)
        pil_img = pil_img.resize((target_width, target_height), Image.Resampling.LANCZOS)
        
        # 保存到内存流
        image_stream = io.BytesIO()
        pil_img.save(image_stream, format='PNG')
        image_stream.seek(0)
        # -------------------------------------------------------------

        # 插入图片 (xlsxwriter 支持直接传入内存流，需提供一个虚拟文件名)
        worksheet.insert_image(row_idx, 0, 'image.png', {
            'image_data': image_stream,
            'object_position': 1  # 1 表示图片随单元格移动和调整大小
        })
        
        # 写入文本描述
        worksheet.write(row_idx, 1, item['label'], center_format)
        
        # 第3列留空并居中
        worksheet.write(row_idx, 2, '', center_format)
        
        # 设置行高为 90 磅 (约 1.25 英寸)，确保能完整容纳 115px 高的图片
        worksheet.set_row(row_idx, 90)

    # 6. 关闭并保存文件
    workbook.close()
    logging.info(f"{output_path} 文件保存成功")
    
    return output_path


if __name__ == "__main__":

    from core.agent import MultClient
    import asyncio
    from dotenv import load_dotenv
    import os
    load_dotenv()
    api_key:str = os.getenv("api_key")
    base_url:str = os.getenv("base_url")
    model_name:str = os.getenv("model")
    example:list[str] = ["未固定的高空作业平台","裸露的带电电缆"]
    image_path = "image/6.jpg"
    ima_list = ['image/6.jpg']

    client = MultClient(api_key,base_url,model_name,example)
    
    async def test() -> None:
        example_data = await MultClient.infer(image_path)
        #print(example_data)
        export_to_word(example_data)

    asyncio.run(test())

    #api = MultClient(api_key,base_url,model_name,example)
    #example_data = api.batch_infer(ima_list,True)
    #export_to_word(example_data)
    
    